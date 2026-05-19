from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fpdf import FPDF


def _snippet(text: str, max_len: int) -> str:
    t = text.strip()
    if len(t) <= max_len:
        return t
    return t[: max_len - 3] + "..."


def export_docx(doc_title: str, qa_pairs: List[Dict], output_path: str) -> str:
    doc = Document()

    title = doc.add_heading("Research Q&A Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Source document: {doc_title}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for i, qa in enumerate(qa_pairs, start=1):
        doc.add_heading(f"Q{i}: {qa['question']}", level=2)

        answer_para = doc.add_paragraph(qa["answer"])
        answer_para.style.font.size = Pt(11)

        if qa.get("citations"):
            doc.add_paragraph("Sources used:", style="Caption")
            for citation in qa["citations"]:
                excerpt = citation.get("preview") or _snippet(citation["text"], 120)
                cite_text = (
                    f"  Page {citation['page']}  —  "
                    f"Confidence: {citation['confidence']}%  —  "
                    f"\"{_snippet(excerpt, 120)}\""
                )
                doc.add_paragraph(cite_text, style="Caption")

        doc.add_paragraph("")

    doc.save(output_path)
    return output_path


def export_pdf(doc_title: str, qa_pairs: List[Dict], output_path: str) -> str:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, "Research Q&A Report", ln=True, align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, f"Source: {doc_title}", ln=True, align="C")
    pdf.ln(8)

    for i, qa in enumerate(qa_pairs, start=1):
        pdf.set_font("Helvetica", "B", 13)
        pdf.multi_cell(0, 8, f"Q{i}: {qa['question']}")
        pdf.ln(2)

        pdf.set_font("Helvetica", "", 11)
        pdf.multi_cell(0, 7, qa["answer"])
        pdf.ln(3)

        if qa.get("citations"):
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(100, 100, 100)
            for citation in qa["citations"]:
                excerpt = citation.get("preview") or _snippet(citation["text"], 100)
                cite_text = (
                    f"Page {citation['page']}  |  "
                    f"Confidence: {citation['confidence']}%  |  "
                    f"\"{_snippet(excerpt, 100)}\""
                )
                pdf.multi_cell(0, 6, cite_text)
            pdf.set_text_color(0, 0, 0)

        pdf.ln(5)

    pdf.output(output_path)
    return output_path
